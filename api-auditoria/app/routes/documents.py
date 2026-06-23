"""
MÓDULO: RUTAS API - CARGA DE DOCUMENTOS
RESPONSABLE: Desarrollador Backend / Pipeline RAG.
"""

from fastapi import APIRouter, UploadFile, File
from typing import List
import shutil
import os
from langchain_community.document_loaders import PyPDFLoader
import traceback

router = APIRouter(prefix="/api/v1/documents", tags=["Documentos"])

@router.post("/upload")
async def upload_pdf(file: UploadFile = File(...)):
    print(f"📥 Recibiendo archivo: {file.filename}...")
    temp_file_path = f"temp_{file.filename}"
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    extracted_text = ""
    try:
        filename_lower = file.filename.lower()
        if filename_lower.endswith(".pdf"):
            print("📖 Leyendo el PDF...")
            loader = PyPDFLoader(temp_file_path)
            documents = loader.load()
            extracted_text = "\n\n".join([doc.page_content for doc in documents])
        elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
            print("📖 Leyendo el Word...")
            try:
                from docx import Document
                doc = Document(temp_file_path)
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e1:
                print(f"⚠️ python-docx falló (probablemente es un .doc antiguo). Intentando docx2txt... Error: {e1}")
                try:
                    import docx2txt
                    extracted_text = docx2txt.process(temp_file_path)
                except Exception as e2:
                    print(f"⚠️ docx2txt falló. Intentando extraer .doc con win32com (MS Word)... Error: {e2}")
                    try:
                        import win32com.client
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        abs_path = os.path.abspath(temp_file_path)
                        word_doc = word.Documents.Open(abs_path)
                        extracted_text = word_doc.Range().Text
                        word_doc.Close(False)
                        word.Quit()
                        print("✅ Texto extraído de .doc antiguo usando win32com.")
                    except Exception as e3:
                        extracted_text = f"Error: No se pudo leer el archivo Word. {e3}"
        else:
            raise ValueError("Formato de archivo no soportado. Solo PDF o Word.")
            
        print("✅ Texto extraído con éxito. Guardando en PostgreSQL...")
        
        from app.database import SessionLocal
        from app.models import DocumentRecord
        
        db = SessionLocal()
        doc = db.query(DocumentRecord).filter(DocumentRecord.filename == file.filename).first()
        if doc:
            doc.content = extracted_text
        else:
            doc = DocumentRecord(filename=file.filename, content=extracted_text)
            db.add(doc)

        db.commit()
        db.close()

        os.remove(temp_file_path)
        
        return {
            "filename": file.filename,
            "status": "Recibido con éxito",
            "extracted_text": f"[PG_DOC_REF:{file.filename}]",
            "detail": f"Documento guardado en base de datos. Tamaño del texto: {len(extracted_text)} caracteres."
        }
    except Exception as e:
        print(f"❌ Error al procesar: {traceback.format_exc()}")
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        return {
            "filename": file.filename,
            "status": "Error",
            "extracted_text": "",
            "detail": str(e)
        }

@router.post("/admin-upload")
async def upload_admin_pdf(files: List[UploadFile] = File(...)):
    print(f"📥 [ADMIN] Recibiendo {len(files)} archivos para ChromaDB...")
    
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from app.rag.vector_store import add_documents_to_chroma
        from langchain_core.documents import Document as LcDocument
        
        total_chunks_generados = 0
        nombres_archivos = []
        
        for file in files:
            temp_file_path = f"temp_admin_{file.filename}"
            with open(temp_file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
                
            extracted_text = ""
            filename_lower = file.filename.lower()
            if filename_lower.endswith(".pdf"):
                loader = PyPDFLoader(temp_file_path)
                documents = loader.load()
            elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(temp_file_path)
                    extracted_text = "\n".join([para.text for para in doc.paragraphs])
                    documents = [LcDocument(page_content=extracted_text, metadata={"source": file.filename})]
                except Exception as e1:
                    print(f"⚠️ python-docx falló en {file.filename} (puede ser un .doc antiguo). Intentando docx2txt... Error: {e1}")
                    try:
                        import docx2txt
                        extracted_text = docx2txt.process(temp_file_path)
                        documents = [LcDocument(page_content=extracted_text, metadata={"source": file.filename})]
                    except Exception as e2:
                        print(f"⚠️ docx2txt falló. Intentando extraer .doc con win32com (MS Word)... Error: {e2}")
                        try:
                            import win32com.client
                            word = win32com.client.Dispatch("Word.Application")
                            word.Visible = False
                            abs_path = os.path.abspath(temp_file_path)
                            word_doc = word.Documents.Open(abs_path)
                            extracted_text = word_doc.Range().Text
                            word_doc.Close(False)
                            word.Quit()
                            documents = [LcDocument(page_content=extracted_text, metadata={"source": file.filename})]
                            print(f"✅ Texto extraído de {file.filename} usando win32com.")
                        except Exception as e3:
                            print(f"❌ Fallo total al leer {file.filename}. Omitiendo. Error: {e3}")
                            os.remove(temp_file_path)
                            continue
            else:
                os.remove(temp_file_path)
                continue # Skip non-supported files
                
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            chunks = text_splitter.split_documents(documents)
            add_documents_to_chroma(chunks)
            total_chunks_generados += len(chunks)
            nombres_archivos.append(file.filename)
            os.remove(temp_file_path)
            
        return {
            "archivos_procesados": nombres_archivos,
            "status": "Archivos vectorizados con éxito en ChromaDB",
            "chunks_generados": total_chunks_generados
        }
    except Exception as e:
        print(f"❌ Error en admin upload múltiple: {traceback.format_exc()}")
        return {"status": "Error", "detail": str(e)}

@router.get("/count")
async def count_documents():
    """Endpoint para contar chunks en ChromaDB"""
    try:
        from app.rag.vector_store import get_vector_store
        vs = get_vector_store()
        # Note: Chroma client collection count
        count = vs._collection.count()
        return {"total_chunks_en_chromadb": count, "message": "Conectado a ChromaDB Server."}
    except Exception as e:
        return {"total_chunks_en_chromadb": 0, "message": f"Error conectando a ChromaDB: {e}"}

@router.get("/vectors")
async def get_vectors_sample():
    """Endpoint para obtener una muestra visual de los vectores para el Admin Frontend"""
    try:
        from app.rag.vector_store import get_vector_store
        vs = get_vector_store()
        collection = vs._collection
        
        if collection.count() == 0:
            return {"status": "empty", "message": "No hay vectores almacenados todavía."}
            
        data = collection.get(include=["embeddings", "documents", "metadatas"])
        
        if not data or not data.get('ids') or len(data['ids']) == 0:
            return {"status": "error", "message": "No se pudo extraer la muestra."}
            
        text = data['documents'][0]
        metadata = data['metadatas'][0]
        
        # Construir matriz visual con todos los números
        matrix_repr = []
        for vector in data['embeddings']:
            first_three = ", ".join([f"{num:.3f}" for num in vector[:3]])
            last_three = ", ".join([f"{num:.3f}" for num in vector[-3:]])
            matrix_repr.append(f"[{first_three}, ... , {last_three}]")
            
        final_matrix = "[" + ", ".join(matrix_repr) + "]"
        
        return {
            "status": "success",
            "dimensiones": len(data['embeddings'][0]),
            "texto": text,
            "archivo": metadata.get('source', 'Desconocido'),
            "matrix_string": final_matrix
        }
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return {"status": "error", "message": str(e)}